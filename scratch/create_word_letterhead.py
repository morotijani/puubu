import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_word_letterhead():
    output_path = r"C:\Users\tijan\.gemini\antigravity\brain\a30eef2a-ac1c-49a5-90e5-6e81a3183b2f\Kokuromotie_Letterhead.docx"
    logo_path = r"c:\xampp\htdocs\kokuromotie\media\logo_transparent.png"
    
    document = Document()
    
    # Set document margins (A4 standard)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(1.5)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
        
    # --- HEADER ---
    header = document.sections[0].header
    
    # Clear any existing paragraphs
    for p in header.paragraphs:
        p.text = ""
        
    # Add a table for layout (Logo left, Text right)
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths (roughly 1.5 inch for logo, 5 for text)
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5.0)
    
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    
    # Insert Logo
    p_left = cell_left.paragraphs[0]
    if os.path.exists(logo_path):
        run = p_left.add_run()
        run.add_picture(logo_path, width=Inches(0.8))
        
    # Insert Company Name
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_brand = p_right.add_run("KOKUROMOTIE\n")
    run_brand.bold = True
    run_brand.font.size = Pt(20)
    run_brand.font.color.rgb = RGBColor(29, 61, 55) # Dark Green
    
    run_slogan = p_right.add_run("Professional Digital Electoral System")
    run_slogan.font.size = Pt(9)
    run_slogan.font.color.rgb = RGBColor(163, 230, 53) # Lime Green
    
    # Add a bottom border to header (by adding a paragraph with a border, or just a line of underscores for simplicity)
    p_line = header.add_paragraph()
    run_line = p_line.add_run("_" * 65)
    run_line.font.color.rgb = RGBColor(29, 61, 55)
    run_line.font.size = Pt(12)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    # --- FOOTER ---
    footer = document.sections[0].footer
    for p in footer.paragraphs:
        p.text = ""
        
    f_line = footer.add_paragraph()
    f_run_line = f_line.add_run("_" * 65)
    f_run_line.font.color.rgb = RGBColor(163, 230, 53)
    f_run_line.font.size = Pt(12)
    f_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    f_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    f_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    f_table.columns[0].width = Inches(2.1)
    f_table.columns[1].width = Inches(2.1)
    f_table.columns[2].width = Inches(2.1)
    
    # Address
    c1 = f_table.cell(0, 0).paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = c1.add_run("ADDRESS\n")
    r1.bold = True
    r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor(29, 61, 55)
    c1.add_run("Kokuromotie Headquarters\nCity, Region, ZIP").font.size = Pt(8)
    
    # Contact
    c2 = f_table.cell(0, 1).paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = c2.add_run("CONTACT\n")
    r2.bold = True
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(29, 61, 55)
    c2.add_run("+1 (555) 123-4567\ninfo@kokuromotie.com").font.size = Pt(8)
    
    # Online
    c3 = f_table.cell(0, 2).paragraphs[0]
    c3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = c3.add_run("ONLINE\n")
    r3.bold = True
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(29, 61, 55)
    c3.add_run("www.kokuromotie.com\n@Kokuromotie").font.size = Pt(8)


    # --- BODY ---
    # Top metadata block
    p_meta = document.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_ref = p_meta.add_run("Ref: ")
    run_ref.bold = True
    p_meta.add_run("KOK-2026-001\n")
    run_date = p_meta.add_run("Date: ")
    run_date.bold = True
    p_meta.add_run("[Current Date]")
    
    document.add_paragraph("\n")
    
    p_to = document.add_paragraph()
    run_to = p_to.add_run("To: ")
    run_to.bold = True
    p_to.add_run("[Recipient Name/Organization]")
    
    document.add_paragraph("\n")
    
    p_title = document.add_paragraph("[Document Title / Subject Line]")
    p_title.style = 'Title'
    for r in p_title.runs:
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(29, 61, 55)
    
    document.add_paragraph("Dear [Name],")
    document.add_paragraph("Start typing your official correspondence here. Because this is a native Microsoft Word document, you can easily type anywhere, format text, and insert long tables without writing any code.")
    document.add_paragraph("The logo, company branding, and contact details are safely stored in the document's Header and Footer. You can double-click the top or bottom of the page to edit those details at any time.")
    document.add_paragraph("Sincerely,\n\n\n\n[Your Name]\nElectoral Systems Director\nKokuromotie")

    document.save(output_path)
    print(f"Native Word document created at: {output_path}")

if __name__ == "__main__":
    create_word_letterhead()
